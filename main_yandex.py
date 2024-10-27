from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uuid
import json
import re
import clickhouse_connect
from docx import Document
from langchain_community.embeddings.yandex import YandexGPTEmbeddings
from vectorestore_yandex import Vectorestore
from fastapi.responses import JSONResponse

app = FastAPI(title="RAG System API", description="API для загрузки и обработки .docx документов", version="1.0.0")

# Настройка CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Разрешить запросы с любых источников
    allow_credentials=True,
    allow_methods=["*"],  # Разрешить все методы (GET, POST и т.д.)
    allow_headers=["*"]   # Разрешить все заголовки
)

# Настройка клиента для подключения к ClickHouse
client = clickhouse_connect.get_client(
    host='0.0.0.0',
    port=8123,
    username='default',
    password=''
)

# Убедитесь, что таблица в ClickHouse существует
client.command("CREATE DATABASE IF NOT EXISTS rag_system")
client.command('''
    CREATE TABLE IF NOT EXISTS rag_system.docs_ya (
        metadata String,
        content String,
        embedding Array(Float32)
    ) ENGINE = MergeTree()
    ORDER BY tuple()
''')

# Настройка эмбеддингов
embeddings = YandexGPTEmbeddings(api_key='AQVNwwO6r_Ko914j7zAW3H5lPga-sVKoXCipbr8_', folder_id='b1gp8gibtjpdhh1rk46d')
vectorestore = Vectorestore(docx_path="./RAG-Documents/data.docx")

class QueryRequest(BaseModel):
    query: str
    chat_id: str
    user_id: str

class QueryResponse(BaseModel):
    answer: str

class CreateChatRequest(BaseModel):
    user_id: str
    document_id: str

class ChatInfo(BaseModel):
    chat_id: str
    user_id: str
    document_id: str

class UploadResponse(BaseModel):
    document_id: str

def extract_chapters_from_docx(doc):
    chapters = []
    chapter_number_regex = re.compile(r'^\d+(\.\d+)*$')
    current_number = ""

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            text = paragraph.text.strip()
            split_text = text.split(" ", 1)
            if len(split_text) > 1 and chapter_number_regex.match(split_text[0]):
                chapters.append(text)
            else:
                full_title = f"{current_number} {text}" if current_number else text
                chapters.append(full_title)
                current_number = ""
        elif chapter_number_regex.match(paragraph.text.strip()):
            current_number = paragraph.text.strip()

    return chapters

def split_document_by_headings(doc, headings):
    content_dict = {}
    current_heading = None
    current_text = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in headings:
            if current_heading and current_text:
                content_dict[current_heading] = "\n".join(current_text)
                current_text = []
            current_heading = text
        elif current_heading:
            current_text.append(text)

    if current_heading and current_text:
        content_dict[current_heading] = "\n".join(current_text)

    return content_dict

def create_rag_documents(content_dict):
    rag_documents = []
    for key, content in content_dict.items():
        embedding = embeddings.embed_query(content)
        document = {
            "metadata": key,
            "content": content,
            "embedding": embedding
        }
        rag_documents.append(document)
    return rag_documents

@app.post("/get_answer", response_model=QueryResponse)
async def get_answer(request: QueryRequest):
    try:
        answer = vectorestore.async_get_answer(query=request.query, user_id=request.user_id, chat_id=request.chat_id)
        return JSONResponse(content={"answer": answer})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при получении ответа: {str(e)}")

@app.get("/get_chats", response_model=list[ChatInfo])
async def get_chats():
    try:
        chats = vectorestore.get_last_messages_by_chat()
        return JSONResponse(content=chats)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при получении чатов: {str(e)}")

@app.post("/create_chat", response_model=ChatInfo)
async def create_chat(request: CreateChatRequest):
    try:
        chat_info = vectorestore.create_chat(user_id=request.user_id, document_id=request.document_id)
        return JSONResponse(content=chat_info)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при создании чата: {str(e)}")

@app.post("/upload_doc", response_model=UploadResponse)
async def upload_doc(file: UploadFile = File(...)):
    try:
        # Считываем загруженный файл .docx
        doc = Document(file.file)

        # Извлекаем главы и содержание
        chapters = extract_chapters_from_docx(doc)
        content_dict = split_document_by_headings(doc, chapters)

        # Создаем RAG-документы с эмбеддингами
        rag_documents = create_rag_documents(content_dict)

        # Создаем уникальный document_id
        document_id = str(uuid.uuid4())

        # Подготавливаем данные для вставки в ClickHouse
        data_to_insert = [
            (entry["metadata"], entry["content"], entry["embedding"])
            for entry in rag_documents
        ]
        
        # Вставка данных в ClickHouse
        client.insert('rag_system.docs_ya', data_to_insert, column_names=["metadata", "content", "embedding"])

        return {"document_id": document_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке документа: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)
